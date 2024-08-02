import os
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize

nltk.download('punkt')  # Download the tokenizer models

def find_word_in_docx(file_path, search_word):
    """Extract sentences containing the search_word from a docx file."""
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return []
    
    sentences_with_word = []
    for para in doc.paragraphs:
        sentences = sent_tokenize(para.text)
        for sentence in sentences:
            if search_word.lower() in sentence.lower():
                sentences_with_word.append(sentence)
    return sentences_with_word

def search_in_folder(folder_path, search_word):
    """Walk through all folders and subfolders to find .docx files and search for the word."""
    results = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.endswith('.docx') and not file.startswith('~$'):
                file_path = os.path.join(root, file)
                sentences = find_word_in_docx(file_path, search_word)
                for sentence in sentences:
                    results.append((file, sentence))
    return results

def write_results_to_docx(results, output_file):
    """Write the results to a new docx file in tabular format."""
    doc = Document()
    doc.add_heading('Search Results', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'File Name'
    hdr_cells[1].text = 'Sentence'

    for file_name, sentence in results:
        row_cells = table.add_row().cells
        row_cells[0].text = file_name
        row_cells[1].text = sentence

    doc.save(output_file)

def main():
    parent_folder = input("Enter the path of the parent folder: ")
    search_word = input("Enter the word to search for: ")
    output_file = "search_results.docx"

    results = search_in_folder(parent_folder, search_word)
    write_results_to_docx(results, output_file)
    print(f"Results have been saved to {output_file}")

if __name__ == "__main__":
    main()
